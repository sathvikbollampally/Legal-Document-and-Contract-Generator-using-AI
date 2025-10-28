from groq import Groq
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# === API Setup ===
from dotenv import load_dotenv
import os
import re
# Load the .env file
load_dotenv()

# Get the API key from environment variable
API = os.getenv("GROQ_API_KEY")

from groq import Groq  # or wherever you import Groq from

client = Groq(api_key=API)


# === Role Mapping ===
def get_roles(contract_type):
    roles = {
        "Construction Agreement": ("Contractor", "Owner"),
        "Rental Agreement": ("Landlord", "Tenant"),
        "Service Contract": ("Service Provider", "Client"),
        "Sale Deed": ("Seller", "Buyer"),
        "Gift Deed": ("Donor", "Donee"),
        "Transfer of Property": ("Transferor", "Transferee")
    }
    return roles.get(contract_type, ("Party A", "Party B"))

# === Prompt Generator ===
def generate_prompt(inputs):
    role_a, role_b = get_roles(inputs['contract_type'])

    return f"""
You are a legal AI that drafts formal contracts. Draft a comprehensive agreement titled:

Start the document with the following paragraph (left-aligned, not centered):

"This {inputs['contract_type']} ('Agreement') is entered into and made effective as of {inputs['effective_date']} ('Effective Date'), by and between:

{inputs['party_a']}, acting as the '{role_a}',
and
{inputs['party_b']}, acting as the '{role_b}'.

The parties hereby agree as follows:"

Then continue with structured legal content using these numbered articles:

- ARTICLE 1: DEFINITIONS  
- ARTICLE 2: OBLIGATIONS OF THE PARTIES  
{"" if inputs['contract_type'] == 'Gift Deed' else "- ARTICLE 3: PAYMENT TERMS"}  
- ARTICLE 4: INDEMNITY  
- ARTICLE 5: TERMINATION  
- ARTICLE 6: FORCE MAJEURE  
- ARTICLE 7: GOVERNING LAW  
- ARTICLE 8: ENTIRE AGREEMENT  
- ARTICLE 9: SIGNATURES

Include the following project/property details:
- Description: {inputs['project_type']}
- Duration: {inputs['duration'] or 'N/A'}
{"- Property Location: " + inputs['property_location'] if inputs['contract_type'] in ['Sale Deed', 'Transfer of Property'] else ""}
{"- Possession Date: " + inputs['possession_date'] if inputs['contract_type'] in ['Sale Deed', 'Transfer of Property'] else ""}
{"- Payment: ₹" + inputs['price'] if inputs['contract_type'] in ['Sale Deed', 'Transfer of Property'] else ""}
{"- Payment Schedule: " + inputs['payment_schedule'] if inputs['contract_type'] in ['Construction Agreement', 'Rental Agreement', 'Service Contract', 'Sale Deed', 'Transfer of Property'] else ""}
{"- Payment Date: " + inputs['payment_date'] if inputs['contract_type'] in ['Sale Deed', 'Transfer of Property'] else ""}
{"- Currency: " + inputs['currency'] if inputs['contract_type'] in ['Sale Deed', 'Transfer of Property'] else ""}
- Governing Law: {inputs['jurisdiction']}

Formatting & Content Requirements:
1. Title must be centered, bold, and styled legally in Times New Roman, 16pt.
2. Start with the provided introductory paragraph, aligned left.
3. Use formal legal language throughout the contract.
4. Numbered articles with headings like "ARTICLE 1: DEFINITIONS".
5. Make all article headings bold, 13pt font, left-aligned.
6. Use Times New Roman font, 12pt, for body text with tight spacing.
7. Add a signature section:

{inputs['party_a']} ({role_a})  
Signature: __________________________  
Date: __________________________  

{inputs['party_b']} ({role_b})  
Signature: __________________________  
Date: __________________________  

8. End with a legal disclaimer: “This document is a template and should be reviewed by a legal professional.”
9. Do not include markdown, page numbers, or version info.
10. Ensure the entire content is output as plain text ready for Word formatting.
""".strip()

# === Groq Contract Text Generator ===
def get_contract_text(prompt):
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a legal contract drafting assistant. Write professional, clearly formatted, "
                        "compliant contracts using formal language and legal structure. Never offer legal advice."
                    )
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Oops.. Error generating contract: {e}"

def string_to_word_doc(text, filename,title=None):
    doc = Document()

    # === Set default font style ===
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # === Clean up text ===
    # Remove markdown bold (**) and HTML-like tags
    cleaned_text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # remove **bold**
    cleaned_text = re.sub(r"<[^>]+>", "", cleaned_text)   # remove HTML tags
    cleaned_text = cleaned_text.strip()

    lines = cleaned_text.split("\n")

    # === Try to detect and format the title ===
    title = None
    for i, line in enumerate(lines):
        if "agreement" in line.lower():
            title = line.strip()
            lines.pop(i)
            break

    if title:
        title_para = doc.add_paragraph()
        run = title_para.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(16)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # add spacing

    # === Process lines ===
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # ARTICLE headings (bold, uppercase)
        if re.match(r"^article\s+\d+", stripped.lower()):
            p = doc.add_paragraph()
            run = p.add_run(stripped.upper())
            run.bold = True
            run.font.size = Pt(13)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            continue

        # Legal disclaimer (bold)
        if "this document is a template and should be reviewed by a legal professional" in stripped.lower():
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)
            continue

        # Regular text
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)

    # === Save the document ===
    doc.save(filename)
    return filename